# coding=utf-8

import docx
import os
import re
#百度通用翻译API,不包含词典、tts语音合成等资源，如有相关需求请联系translate_api@baidu.com
import http.client
import hashlib
import urllib
import random
import json

appid = ''  # 填写你的appid
secretKey = ''  # 填写你的密钥

httpClient = None
myurl = '/api/trans/vip/translate'

fromLang = 'en'   #原文语种
toLang = 'zh'   #译文语种
salt = random.randint(32768, 65536)

#txt文本路径
txts = os.listdir('F:/mc_learn/pdf/test/txt/59_pdf_output_dir')
file = docx.Document()
much = len(txts)
print(much)
num = 0
for i in range(1,much+1):
	file_object = open('F:/mc_learn/pdf/test/txt/59_pdf_output_dir/pages-{}-{}.txt'.format(i,i),'r', encoding='utf-8')
	filelines = file_object.readlines()
	paragraph = ""
	try: 
	    for line in filelines:
	    	line = line.strip()
	    	line = line.strip('\n')
	    	paragraph = paragraph + line
	    	# print(line)
	finally:
		file_object.close()

	q = paragraph
	sign = appid + q + str(salt) + secretKey
	sign = hashlib.md5(sign.encode()).hexdigest()
	myurl = myurl + '?appid=' + appid + '&q=' + urllib.parse.quote(q) + '&from=' + fromLang + '&to=' + toLang + '&salt=' + str(
	salt) + '&sign=' + sign

	try:
	    httpClient = http.client.HTTPConnection('api.fanyi.baidu.com')
	    httpClient.request('GET', myurl)

	    # response是HTTPResponse对象
	    response = httpClient.getresponse()
	    result_all = response.read().decode("utf-8")
	    result = json.loads(result_all)

	    result = result['trans_result']
	    result = result[0]['dst']

	except Exception as e:
	    print (e)
	finally:
	    if httpClient:
	        httpClient.close()

	file.add_paragraph(u'{}'.format(result))
	file.add_paragraph('【{}】'.format(num))
	num = num + 1


file.save('F:/mc_learn/pdf/test/{}_ch.docx'.format(num))