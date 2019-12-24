#encoding=utf-8
import docx
import os
import re

#txt文本合并至word
txts = os.listdir('F:/mc_learn/pdf/test/txt/59_pdf_output_dir')
file = docx.Document()
much = len(txts)
print(much)
num = 0
for i in range(1,much+1):
	file_object = open('F:/mc_learn/pdf/test/txt/59_pdf_output_dir/pages-{}-{}.txt'.format(i,i),'r', encoding='utf-8')
	filelines = file_object.readlines()
	paragraph = ""
	try: 
	    for line in filelines:
	    	line = line.strip()
	    	line = line.strip('\n')
	    	paragraph = paragraph + line
	    	# print(line)
	finally:
		file_object.close()

	file.add_paragraph(u'{}'.format(paragraph))
	file.add_paragraph('【{}】'.format(num))
	num = num + 1


file.save('F:/mc_learn/pdf/test/{}.docx'.format(num))